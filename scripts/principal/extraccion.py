# -*- coding: utf-8 -*-
import os,sys
reload(sys)
sys.setdefaultencoding('utf8')
import codecs
import docx
import re
import unicodedata
global jsonDatos, matriz, last_sequence
matriz=[]

pat = re.compile(r'\s+')
flag_acumular_diagnostico=False
count_lineas=0
jsonDatos= {'id_tblpat':"",
			'nobio':"",
			'fechabx':"",
			'cedula':"",
			'ape1':"",
			'ape2':"",
			'apec':"",
			'nom1':"",
			'nom2':"",
			'labbio':"",
			'recolector_':"",
			'link':"",
			'datesys':"",
			'fuente_recoleccion':"",
			'estado':"",
			'bdrpcc':"",
			'anexado':"",
			'recolector':"",
			'sexo':"",
			'edad':"",
			'cedula1':"",
			'cedula2':"",
			'fechainfo':"",
			'entidad':"",
			'diagnostico':"",
			'medico':"",
			'desc_macro':"",
			'desc_micro':""
			}



#Clase para leer el html, se obtienen solo el Data, se limpia y se arma la lista.
class LecturaDOCX():
	

	def inicializar(self):
		global jsonDatos		
		count_lineas=0
		jsonDatos= {'id_tblpat':"",
				'nobio':"",
				'fechabx':"",
				'cedula':"",
				'ape1':"",
				'ape2':"",
				'apec':"",
				'nom1':"",
				'nom2':"",
				'labbio':"",
				'recolector_':"",
				'link':"",
				'datesys':"",
				'fuente_recoleccion':"",
				'estado':"",
				'bdrpcc':"",
				'anexado':"",
				'recolector':"",
				'sexo':"",
				'edad':"",
				'cedula1':"",
				'cedula2':"",
				'fechainfo':"",
				'entidad':"",
				'diagnostico':"",
				'medico':"",
				'desc_macro':"",
				'desc_micro':""
				}

	def crearMatriz(self):		
		global matriz
		matriz=[[None] * 28 for i in range(1)]
		matriz[0][0]="id_tblpat"
		matriz[0][1]="nobio"
		matriz[0][2]="fechabx"
		matriz[0][3]="cedula"
		matriz[0][4]="ape1"
		matriz[0][5]="ape2"
		matriz[0][6]="apec"
		matriz[0][7]="nom1"
		matriz[0][8]="nom2"
		matriz[0][9]="labbio"
		matriz[0][10]="recolector_"
		matriz[0][11]="link"
		matriz[0][12]="datesys"
		matriz[0][13]="fuente_recoleccion"
		matriz[0][14]="estado"
		matriz[0][15]="bdrpcc"
		matriz[0][16]="anexado"
		matriz[0][17]="recolector"
		matriz[0][18]="sexo"
		matriz[0][19]="edad"
		matriz[0][20]="cedula1"
		matriz[0][21]="cedula2"
		matriz[0][22]="fechainfo"
		matriz[0][23]="entidad"
		matriz[0][24]="diagnostico"
		matriz[0][25]="medico"
		matriz[0][26]="desc_macro"
		matriz[0][27]="desc_micro"
		# print matriz
		self.crearFila()

	def crearFila(self):
		global matriz, jsonDatos
		if jsonDatos["nobio"]!='' or jsonDatos["cedula"]!='':
			fila=[]
			for item in matriz[0]:
				fila.append(jsonDatos[item])		
			matriz.append(fila)

	def extraerEntidades(self,ruta_archivo, val_sequence):
		global jsonDatos, last_sequence
		last_sequence=val_sequence
		doc = docx.Document(ruta_archivo)
		fullText = []
		nombres=[]
		count_lineas=0
		numero_parrafos=len(doc.paragraphs)
		self.crearMatriz()
		for i in range(0,numero_parrafos):
			try:
				if doc.paragraphs[i].text=="":
					# print "**",jsonDatos				
					last_sequence+=1
					self.crearFila()
					self.inicializar()
					count_lineas=0

				elif(count_lineas<4):
					count_lineas+=1
					linea=doc.paragraphs[i].text.strip().split(":")
					if linea[0].strip()=="NOMBRE":
						jsonDatos['id_tblpat']=str(last_sequence)
						self.extraerNombre_Documento(linea)

					elif linea[0].strip()=="REF.":
						self.extraer_ref_edad_sexo_codigo(linea)

					elif linea[0].strip()=="MEDICO":
						self.extraer_medico_fechrecibido(linea)

					elif linea[0].strip()=="ENTIDAD":
						count_lineas=4
						self.extraer_entidad_fechinfo(linea)

					else:
						count_lineas=4
						self.extraer_diagnostico(doc.paragraphs[i].text.strip())


				elif (count_lineas==4):
					self.extraer_diagnostico(doc.paragraphs[i].text.strip())
		
			except IndexError as e:
				print "Error en WORD: buscar la palabra: ", doc.paragraphs[i].text


			

	def extraerNombre_Documento(self, listanombre):
		lista_palabras=pat.sub(' ', listanombre[1]).split(" ")
		if len(lista_palabras)==5:
			jsonDatos['nom1']=lista_palabras[0]
			jsonDatos['nom2']=lista_palabras[1]
			jsonDatos['ape1']=lista_palabras[2]
			jsonDatos['ape2']=lista_palabras[3]
		elif len(lista_palabras)==4:
			jsonDatos['nom1']=lista_palabras[0]
			jsonDatos['ape1']=lista_palabras[1]
			jsonDatos['ape2']=lista_palabras[2]
		elif len(lista_palabras)==3:
			jsonDatos['nom1']=lista_palabras[0]
			jsonDatos['ape1']=lista_palabras[1]
		elif len(lista_palabras)>5:
			jsonDatos['nom1']=lista_palabras[0]
			jsonDatos['nom2']=lista_palabras[1]
			jsonDatos['ape1']=lista_palabras[2]
			jsonDatos['apec']=lista_palabras[3]
			jsonDatos['ape2']=lista_palabras[4]
		if len(listanombre)>2:
			jsonDatos['cedula']=listanombre[2]
		

	def extraer_ref_edad_sexo_codigo(self, listaref):
		nob=pat.sub(' ', listaref[1].strip()).split(" ")
		eda=pat.sub(' ', listaref[2].strip()).split(" ")
		sex=pat.sub(' ', listaref[3].strip()).split(" ")

		if len(nob)>1:
			jsonDatos['nobio']=pat.sub(' ', listaref[1].strip()).split(" ")[0]
		
		if len(eda)>1:
			jsonDatos['edad']=pat.sub(' ', listaref[2].strip()).split(" ")[0]
		
		sexo=pat.sub(' ', listaref[3].strip()).split(" ")[0]		
		if sexo=="M":
			jsonDatos['sexo']=str(1)
		else:
			jsonDatos['sexo']=str(2)

		
	def extraer_medico_fechrecibido(self, listamedico):
		if len(listamedico)==3:
			jsonDatos['medico']=' '.join(pat.sub(' ', listamedico[1].strip()).split(" ")[:-1])
			jsonDatos['fechabx']=pat.sub(' ', listamedico[2].strip()).split(" ")[0]
		if len(listamedico)==2:
			medico=listamedico[1].strip().split(".")[0][:-1]
			fecha=listamedico[1].strip().split(".")[1][:-1]
			jsonDatos['medico']=' '.join(pat.sub(' ', medico).split(" ")[:-1])
			jsonDatos['fechabx']=fecha.strip()
			print jsonDatos['medico']
			print jsonDatos['fechabx']

	
	def extraer_entidad_fechinfo(self, listaentidad):
		tama=len(listaentidad)
		if tama==3:		
			jsonDatos['entidad']=pat.sub(' ', listaentidad[1].strip()).replace('FECH.INFO','').replace('FECHA','').strip()
			jsonDatos['fechainfo']=pat.sub(' ', listaentidad[2].strip()).split(" ")[0]

	def extraer_diagnostico(self, listadiag):
		jsonDatos['diagnostico']+=self.elimina_tildes(str(listadiag).strip())

	def getMatriz(self):
		return matriz

	def getUpdateValSequence(self):
		global last_sequence
		return last_sequence

	def elimina_tildes(self, cadena):
		s = ''.join((c for c in unicodedata.normalize('NFD',unicode(cadena)) if unicodedata.category(c) != 'Mn'))
		return s.decode()

	############################# EXTRACION PARA CITOPATOLOGIA ################################
	def extraerEntidadesCito(self,ruta_archivo, val_sequence):
		global jsonDatos, last_sequence
		last_sequence=val_sequence
		doc = docx.Document(ruta_archivo)
		numero_parrafos=len(doc.paragraphs)
		self.crearMatriz()
		bandera=False
		for i in range(0,numero_parrafos):
			try:
				linea=doc.paragraphs[i].text
				# print "==>",linea
				if linea.strip()=="" or i==numero_parrafos-1:
					# print "\n**",jsonDatos				
					last_sequence+=1
					self.crearFila()
					self.inicializar()
					bandera=False

				elif(bandera):
					jsonDatos['diagnostico']+=linea.strip()

				else:
					dupla=linea.strip().split(":")
					llave=self.elimina_tildes(dupla[0].strip().lower())
					valor=dupla[1].strip()
					if(llave=="documento" or llave=="cc"):
						jsonDatos['id_tblpat']=str(last_sequence)
						jsonDatos['cedula']=valor
					
					elif(llave=="genero"):
						sexo=valor[0]
						if sexo=="M":
							jsonDatos['sexo']=str(1)
						else:
							jsonDatos['sexo']=str(2)
					
					elif(llave=="fecha de salida"):
						jsonDatos['fechabx']=valor
					
					elif(llave=="paciente"):
						lista_palabras=pat.sub(' ', valor.strip()).strip().split(" ")
						#print lista_palabras
						if len(lista_palabras)==4:
							jsonDatos['nom1']=lista_palabras[0]
							jsonDatos['nom2']=lista_palabras[1]
							jsonDatos['ape1']=lista_palabras[2]
							jsonDatos['ape2']=lista_palabras[3]
						elif len(lista_palabras)==3:
							jsonDatos['nom1']=lista_palabras[0]
							jsonDatos['ape1']=lista_palabras[1]
							jsonDatos['ape2']=lista_palabras[2]
						elif len(lista_palabras)==2:
							jsonDatos['nom1']=lista_palabras[0]
							jsonDatos['ape1']=lista_palabras[1]
						elif len(lista_palabras)>4:
							jsonDatos['nom1']=lista_palabras[0]
							jsonDatos['nom2']=lista_palabras[1]
							jsonDatos['ape1']=lista_palabras[2]
							jsonDatos['apec']=lista_palabras[3]
							jsonDatos['ape2']=lista_palabras[4]
					
					elif(llave=="edad"):
						jsonDatos['edad']=valor
					
					elif(llave=="empresa"):
						jsonDatos['entidad']=valor
					
					elif(llave=="peticion no"):
						jsonDatos['nobio']=valor
						bandera=True
					
					elif(llave=="descripcion macroscopica"):
						jsonDatos['desc_macro']=' '.join(dupla[1:])
					
					elif(llave=="descripcion microscopica"):
						jsonDatos['desc_micro']=' '.join(dupla[1:])
					
					elif(llave=="comentario"):
						jsonDatos['desc_macro']+=" "+valor
					
					elif(llave=="diagnostico"):
						# jsonDatos['diagnostico']=' '.join(dupla[1:])
						jsonDatos['diagnostico']=linea.replace("DIAGNOSTICO:", '').replace("diagnostico:", '').strip()

					elif(bandera):
						jsonDatos['diagnostico']+=linea.strip()


			except IndexError as e:
				print "Error en WORD: buscar la palabra: ", doc.paragraphs[i].text




if __name__ == '__main__':
	None